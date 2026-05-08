#!/usr/bin/env python3
"""Apply Wuhan University of Technology style settings to a generated DOCX."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


CITATION_RE = re.compile(r"(\[\d+(?:[-,，]\d+)*\])")


def set_run_fonts(run, east_asia: str, latin: str, size_pt: float, *, bold: bool = False, superscript: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = latin
    run.font.superscript = superscript
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)


def set_outline_level(paragraph, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level))


def clear_paragraph_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def apply_paragraph_base(paragraph, *, alignment=None, first_indent=0, line_exact=True, before=0, after=0) -> None:
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.first_line_indent = Pt(first_indent)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    if line_exact:
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(20)


def rebuild_with_citation_superscript(paragraph, *, east_asia="宋体", latin="Times New Roman", size=12, bold=False) -> None:
    text = paragraph.text
    if not CITATION_RE.search(text):
        return
    clear_paragraph_runs(paragraph)
    for part in CITATION_RE.split(text):
        if not part:
            continue
        is_citation = bool(CITATION_RE.fullmatch(part))
        run = paragraph.add_run(part)
        set_run_fonts(run, east_asia, latin, size, bold=bold, superscript=is_citation)


def style_paragraphs(doc: Document) -> None:
    in_references = False
    in_english_abstract = False
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        normalized = text.replace(" ", "")
        if normalized == "参考文献":
            in_references = True
        if normalized == "致谢":
            in_references = False

        if index == 0:
            apply_paragraph_base(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0, before=0, after=6)
            for run in paragraph.runs:
                set_run_fonts(run, "宋体", "Times New Roman", 26, bold=True)
            continue

        if text.startswith("论文题目："):
            apply_paragraph_base(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0, before=12, after=12)
            for run in paragraph.runs:
                set_run_fonts(run, "黑体", "Times New Roman", 22, bold=True)
            continue

        if text.startswith(("学", "专业班级：", "学生姓名：", "指导教师：")) and "待填写" in text:
            apply_paragraph_base(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0, before=3, after=3)
            for run in paragraph.runs:
                set_run_fonts(run, "宋体", "Times New Roman", 16)
            continue

        if normalized in {"摘要", "目录", "参考文献", "致谢", "学位论文原创性声明"}:
            apply_paragraph_base(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0, before=6, after=6)
            for run in paragraph.runs:
                set_run_fonts(run, "黑体", "Times New Roman", 18, bold=True)
            if normalized in {"参考文献", "致谢"}:
                set_outline_level(paragraph, 0)
            continue

        if normalized.lower() == "abstract":
            apply_paragraph_base(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0, before=6, after=6)
            for run in paragraph.runs:
                set_run_fonts(run, "Times New Roman", "Times New Roman", 18, bold=True)
            in_english_abstract = True
            continue

        if re.match(r"^第\d+章", text):
            apply_paragraph_base(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0, before=6, after=6)
            for run in paragraph.runs:
                set_run_fonts(run, "黑体", "Times New Roman", 18, bold=True)
            set_outline_level(paragraph, 0)
            continue

        if re.match(r"^\d+\.\d+\.\d+", text):
            apply_paragraph_base(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_indent=0, before=6, after=6)
            for run in paragraph.runs:
                set_run_fonts(run, "黑体", "Times New Roman", 14)
            set_outline_level(paragraph, 2)
            continue

        if re.match(r"^\d+\.\d+", text):
            apply_paragraph_base(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_indent=0, before=6, after=6)
            for run in paragraph.runs:
                set_run_fonts(run, "黑体", "Times New Roman", 16)
            set_outline_level(paragraph, 1)
            continue

        if re.match(r"^图\s*\d+[-.]\d+", text) or re.match(r"^表\s*\d+[-.]\d+", text):
            apply_paragraph_base(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0, before=3, after=3)
            for run in paragraph.runs:
                set_run_fonts(run, "宋体", "Times New Roman", 12)
            continue

        if in_references and re.match(r"^\[\d+\]", text):
            apply_paragraph_base(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_indent=-21, before=0, after=0)
            paragraph.paragraph_format.left_indent = Pt(21)
            for run in paragraph.runs:
                set_run_fonts(run, "宋体", "Times New Roman", 10.5)
            continue

        if normalized.startswith("关键词"):
            apply_paragraph_base(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_indent=0, before=0, after=0)
            for run in paragraph.runs:
                set_run_fonts(run, "宋体", "Times New Roman", 12)
            continue

        if text.startswith("Keywords"):
            if text.startswith("Keywords:") and not text.startswith("Keywords: "):
                paragraph.text = text.replace("Keywords:", "Keywords: ", 1)
            apply_paragraph_base(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_indent=0, before=0, after=0)
            for run in paragraph.runs:
                set_run_fonts(run, "Times New Roman", "Times New Roman", 12)
            in_english_abstract = False
            continue

        if in_english_abstract:
            apply_paragraph_base(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_indent=21, before=0, after=0)
            for run in paragraph.runs:
                set_run_fonts(run, "Times New Roman", "Times New Roman", 12)
            continue

        apply_paragraph_base(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_indent=24, before=0, after=0)
        rebuild_with_citation_superscript(paragraph, size=12)
        for run in paragraph.runs:
            if run.font.superscript:
                continue
            set_run_fonts(run, "宋体", "Times New Roman", 12)


def add_field_run(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(end)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def insert_toc_field(doc: Document) -> None:
    paragraphs = list(doc.paragraphs)
    toc_idx = next((i for i, p in enumerate(paragraphs) if p.text.strip() == "目录"), None)
    chapter_indices = [i for i, p in enumerate(paragraphs) if p.text.strip().startswith("第1章")]
    # The markdown source contains a simple visible TOC before the real body.  Use
    # the second "第1章" as the body start when both are present.
    first_chapter_idx = chapter_indices[1] if len(chapter_indices) > 1 else (chapter_indices[0] if chapter_indices else None)
    if toc_idx is None or first_chapter_idx is None or first_chapter_idx <= toc_idx:
        return

    for paragraph in paragraphs[toc_idx + 1 : first_chapter_idx]:
        remove_paragraph(paragraph)

    toc_para = doc.paragraphs[toc_idx]
    new_p = OxmlElement("w:p")
    toc_para._p.addnext(new_p)
    field_para = doc.paragraphs[toc_idx + 1]
    field_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_field_run(field_para, 'TOC \\o "1-3" \\h \\z \\u')
    helper = field_para.add_run("（在 Word 中更新域后生成目录页码）")
    set_run_fonts(helper, "宋体", "Times New Roman", 12)


def setup_sections(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(2.6)
        section.footer_distance = Cm(2.4)

        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = "武汉理工大学本科毕业设计（论文）"
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hp.runs:
            set_run_fonts(run, "宋体", "Times New Roman", 10.5)

        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        clear_paragraph_runs(fp)
        add_field_run(fp, "PAGE")
        for run in fp.runs:
            set_run_fonts(run, "Times New Roman", "Times New Roman", 10.5)


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: format_whut_docx.py <docx>")
        return 1
    path = Path(sys.argv[1])
    doc = Document(str(path))
    setup_sections(doc)
    style_paragraphs(doc)
    insert_toc_field(doc)
    doc.save(str(path))
    print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
